from docx import Document 
import pathlib
import os
import glob
import zipfile
from flask import send_file,Flask,send_from_directory

current_path=pathlib.Path(__file__).parent.resolve()

class SplitDocument:

    def __init__(self,file,filename):
        self.document = Document(file)
        self.filename=filename

    def cleanup(self):
        folder = os.path.join(current_path,"tmp","split","output\*")
        for item in glob.glob(folder):
            if "Blank.zip" not in item:
                #print(item)
                os.remove(os.path.join(folder, item))
    def create_split_sequence(self):
        sections=self.document.sections
        count=0
        master_dict={} 
        self.tr_levels=[]
        for index,paragraph in enumerate (self.document.paragraphs):
            result=paragraph.text.replace(" ","")
            if result:
                self.tr_levels.append({index:paragraph.style.name})
        self.pairs=[]
        pair=[]
        for index, object in enumerate (self.tr_levels):
            key=list(object.keys())[0]
            print(object[key])
            if index > 0 and "heading" in object[key].lower():
                self.pairs.append (pair)
                pair=[]
                pair.append (key)
            elif object==self.tr_levels[-1]:
                pair.append (key)
                self.pairs.append (pair)
            else:
                pair.append (key)
    
    def get_style(self,tr_levels,index):
        for level in tr_levels:
            key=list(level.keys())[0]
            if key==index:
                return level [key]


    def split_documents(self):
        blank_file=os.path.join(current_path, "tmp", "Blank.docx")
        for index, pair in enumerate (self.pairs) :
            doc= Document (blank_file)
            for index,data in enumerate(pair):
                #print(data,document.paragraphs[data].text)
                doc.add_paragraph(self.document.paragraphs [data]. text, style=self.get_style(self.tr_levels,data)) 
                split_file=os.path.join(current_path, "tmp","split","output", f"{self.filename}_{index}.docx")
                doc.save(split_file)

    def download_all(self,filename):
        # Zip file Initialization and you can change the compression type
        rfilename=f'Blank.zip'
        zipfolder = zipfile.ZipFile(rfilename,'w', compression = zipfile.ZIP_STORED)
        #zfile = zipfile.ZipFile(rfilename, "r")

        # zip all the files which are inside in the folder
        file_path=os.path.join(current_path,"tmp","split","output/")
        for root,dirs, files in os.walk(file_path):
            for file in files:
                if file.endswith(".docx"):
                    zfile_path=os.path.join(current_path,"tmp","split","output",file)
                    zipfolder.write(zfile_path)
        zipfolder.close()
        try:
            os.remove( filename + "zip")
        except :
            print("--")

        os.rename(rfilename, filename + "zip") 
        rfilename= filename + "zip"

        return send_file(rfilename,
                mimetype = 'zip',
                as_attachment = True)