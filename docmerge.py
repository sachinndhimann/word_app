import os
from docxcompose.composer import Composer
from docx import Document
import pathlib
import os
import glob
from flask import send_file,Flask,send_from_directory



current_path=pathlib.Path(__file__).parent.resolve()
class DocMerge():
    def __init__(self,list_of_documents,filename):
        self.documents=list_of_documents
        self.filename=filename
        self.merged_doc=Document(os.path.join(current_path,"tmp","Blank.docx"))
    
    def cleanup(self):
        folder = os.path.join(current_path,"tmp","merge","output\*")
        for item in glob.glob(folder):
                os.remove(os.path.join(folder, item))

    def merge(self):
        for path in self.documents:
            doc= Document(path)
            result=doc.paragraphs 
            for value in result:
                self.merged_doc.add_paragraph(value.text,style=value.style.name)
            self.output_path=os.path.join(current_path,"tmp","merge","output","{}_Merged.docx".format(self.filename))
            self.merged_doc.save(self.output_path)
        return self.output_path

    def mergeByPage(self):
        #print(self.documents)
        master = Document(self.documents[0])
        composer = Composer(master)
        self.documents.pop(0)
        #print("master ",master)
        for path in self.documents:
            doc1 = Document(path)
            composer.append(doc1)
        self.output_path=os.path.join(current_path,"tmp","merge","output","{}_Merged.docx".format(self.filename))
        composer.save(self.output_path)
        return self.output_path

        


#obj=DocMerge([r"C:\Users\sachi\Downloads\split_0.docx",r"C:\Users\sachi\Downloads\split_1.docx"],"legal")
#obj.merge()