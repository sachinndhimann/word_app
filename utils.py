from docx import Document

class CreateDocument ():
    def __init__(self, file=None) -> None:
        self.document - Document(file)

    def add_paragraph(self, content, style):
        self.document.add_paragraph(content,style=style)


    def add_heading(self, heading_title, level):
        self.document.add_heading(heading_title, level=level)

    